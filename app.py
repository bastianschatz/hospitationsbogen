# app.py
# Streamlit-Tool: BLI 3.0-basierter Hospitationsbogen (M1–M4) mit Auto-Bewertung & DOCX/JSON/PDF-Export
# Run lokal/Cloud: streamlit run app.py
# Dependencies: streamlit, python-docx, fpdf2

import io
import json
from datetime import datetime
from dataclasses import dataclass, field
from typing import Dict, List

import streamlit as st
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF  # für PDF-Export

st.set_page_config(page_title="Hospitationsbogen (BLI 3.0)", layout="wide")

# ----------------------------- Datenbasis -----------------------------
# Kurzformulierungen entlang BLI 3.0 (Unterrichts-Merkmale M1–M4).
# Schul-intern anpassbar/erweiterbar (z. B. Indikatoren hinzufügen).
BLI_DATA = {
    "M1": {
        "title": "Schülerinnen und Schüler aktivieren",
        "criteria": {
            "1.1": "Kompetenzziele sind für Lernende transparent.",
            "1.2": "Lehrkraft ist sprachbildend (Vorbildfunktion Deutsch).",
            "1.3": "Aktive Beteiligung der Lernenden wird gefördert.",
            "1.4": "Unterricht unterstützt selbstständiges Lernen.",
            "1.5": "Reflexion der Lernprozesse wird angeleitet."
        }
    },
    "M2": {
        "title": "Kompetenzen entwickeln",
        "criteria": {
            "2.1": "Fachlicher Kompetenzzuwachs wird ermöglicht.",
            "2.2": "Medienkompetenz wird gefördert (zielgerichteter Medieneinsatz).",
            "2.3": "Methodenkompetenz wird aufgebaut und angewandt.",
            "2.4": "Deutschkompetenz wird gezielt entwickelt.",
            "2.5": "Fachsprache im DFU wird funktional genutzt."
        }
    },
    "M3": {
        "title": "Unterricht lernwirksam gestalten",
        "criteria": {
            "3.1": "Stundenablauf ist transparent und klar strukturiert.",
            "3.2": "Medien/Arbeitsmittel werden zielgerichtet eingesetzt.",
            "3.3": "Lehrkraft moderiert und steuert Lernprozesse.",
            "3.4": "Heterogenität wird didaktisch berücksichtigt.",
            "3.5": "Personalisiertes/individualisiertes Lernen wird gefördert."
        }
    },
    "M4": {
        "title": "Lernklima förderlich gestalten",
        "criteria": {
            "4.1": "Sozial kompetentes, wertschätzendes Miteinander.",
            "4.2": "Kooperative Lernarrangements unterstützen Soziallernen.",
            "4.3": "Differenzierte, kriteriengeleitete Rückmeldungen.",
            "4.4": "Positive Fehlerkultur ist sichtbar.",
            "4.5": "Lernumgebung unterstützt Lernaktivitäten."
        }
    }
}

RATING_LABELS = {
    0: "0 – nicht beobachtbar",
    1: "1 – Ansatzweise",
    2: "2 – Grundlegend",
    3: "3 – Gut umgesetzt",
    4: "4 – Sehr stark"
}

AUTO_COMMENTS = {
    0: "Bei der Hospitation war dieses Kriterium nicht erkennbar. Mögliche Ursache: Situations-/Phasenabhängigkeit.",
    1: "Ansatzpunkte sind erkennbar. Eine Fokussierung auf klare Routinen/Transparenz könnte die Wirksamkeit erhöhen.",
    2: "Grundlegend vorhanden. Durch Verbindlichkeit/Beispiele/Visualisierung weiter stärken.",
    3: "Überwiegend gut umgesetzt. Punktuell lässt sich die Wirkung noch durch Schüleraktivierung vertiefen.",
    4: "Sehr überzeugend umgesetzt; dient als Good-Practice-Beispiel."
}

# ----------------------------- Profile -----------------------------
# Pro Kolleg*in Fokus-Merkmale und Gewichtungen hinterlegen:
DEFAULT_PROFILES = {
    "— Neu —": {"focus": ["M1", "M3"], "weights": {"M1": 1.0, "M2": 1.0, "M3": 1.2, "M4": 1.0}},
    "Beispiel: Frau Müller": {"focus": ["M2"], "weights": {"M1": 1.0, "M2": 1.3, "M3": 1.0, "M4": 1.0}},
    "Beispiel: Herr Schmidt": {"focus": ["M1", "M4"], "weights": {"M1": 1.2, "M2": 1.0, "M3": 1.0, "M4": 1.2}},
}

if "profiles" not in st.session_state:
    st.session_state["profiles"] = DEFAULT_PROFILES.copy()

# ----------------------------- Dataclasses -----------------------------
@dataclass
class CriterionResult:
    rating: int = 0
    comment: str = ""

@dataclass
class ModuleResult:
    module_key: str = ""
    criteria: Dict[str, CriterionResult] = field(default_factory=dict)

@dataclass
class ObservationForm:
    date: str = ""
    colleague: str = ""
    subject: str = ""
    grade: str = ""
    topic: str = ""
    observer: str = ""
    school: str = ""
    modules: Dict[str, ModuleResult] = field(default_factory=dict)
    strengths: str = ""
    next_steps: str = ""
    profile_focus: List[str] = field(default_factory=list)
    weights: Dict[str, float] = field(default_factory=dict)

# ----------------------------- Helper -----------------------------
def init_form(selected_modules: List[str]) -> ObservationForm:
    form = ObservationForm(
        date=datetime.today().strftime("%Y-%m-%d"),
        modules={}
    )
    for mk in selected_modules:
        mod = ModuleResult(module_key=mk, criteria={})
        for ck in BLI_DATA[mk]["criteria"].keys():
            mod.criteria[ck] = CriterionResult()
        form.modules[mk] = mod
    return form

def compute_scores(form: ObservationForm):
    per_module = {}
    weighted_sum, weight_total = 0.0, 0.0
    for mk, mod in form.modules.items():
        ratings = [c.rating for c in mod.criteria.values()]
        avg = sum(ratings) / len(ratings) if ratings else 0.0
        per_module[mk] = avg
        w = form.weights.get(mk, 1.0)
        weighted_sum += avg * w
        weight_total += w
    overall = (weighted_sum / weight_total) if weight_total else 0.0
    return per_module, overall

def export_docx(form: ObservationForm) -> bytes:
    doc = Document()

    # Styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Header
    h = doc.add_heading('Hospitationsbogen – BLI 3.0', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.add_run('Datum: ').bold = True
    meta.add_run(form.date + '    ')
    meta.add_run('Kolleg*in: ').bold = True
    meta.add_run(form.colleague + '    ')
    meta.add_run('Beobachter*in: ').bold = True
    meta.add_run(form.observer)

    meta = doc.add_paragraph()
    meta.add_run('Fach/Klasse/Thema: ').bold = True
    meta.add_run(f'{form.subject} / {form.grade} / {form.topic}')

    if form.school:
        meta = doc.add_paragraph()
        meta.add_run('Schule: ').bold = True
        meta.add_run(form.school)

    if form.profile_focus:
        pf = doc.add_paragraph()
        pf.add_run('Profil-Fokus: ').bold = True
        pf.add_run(', '.join(form.profile_focus))

    # Tabelle je Modul
    for mk, mod in form.modules.items():
        doc.add_heading(f'{mk} – {BLI_DATA[mk]["title"]}', level=2)
        table = doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Kriterium'
        hdr_cells[1].text = 'Bewertung (0–4)'
        hdr_cells[2].text = 'Kommentar/Hinweis'

        for ck, cres in mod.criteria.items():
            row_cells = table.add_row().cells
            row_cells[0].text = f'{ck} {BLI_DATA[mk]["criteria"][ck]}'
            row_cells[1].text = str(cres.rating)
            row_cells[2].text = cres.comment or ''

        doc.add_paragraph('')

    # Stärken / nächste Schritte
    doc.add_heading('Stärken', level=2)
    doc.add_paragraph(form.strengths or '-')
    doc.add_heading('Nächste Schritte (konkret, terminiert)', level=2)
    doc.add_paragraph(form.next_steps or '-')

    # Scores
    per_module, overall = compute_scores(form)
    doc.add_heading('Zusammenfassung (Scores)', level=2)
    for mk, sc in per_module.items():
        doc.add_paragraph(f'{mk}: {sc:.2f} / 4')
    doc.add_paragraph(f'Gesamt (gewichtet): {overall:.2f} / 4')

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def export_json(form: ObservationForm) -> bytes:
    data = {
        "date": form.date,
        "colleague": form.colleague,
        "subject": form.subject,
        "grade": form.grade,
        "topic": form.topic,
        "observer": form.observer,
        "school": form.school,
        "profile_focus": form.profile_focus,
        "weights": form.weights,
        "modules": {
            mk: {
                "title": BLI_DATA[mk]["title"],
                "criteria": {
                    ck: {
                        "text": BLI_DATA[mk]["criteria"][ck],
                        "rating": cres.rating,
                        "comment": cres.comment
                    } for ck, cres in mod.criteria.items()
                }
            } for mk, mod in form.modules.items()
        },
        "strengths": form.strengths,
        "next_steps": form.next_steps
    }
    return json.dumps(data, ensure_ascii=False, indent=2).encode("utf-8")

def export_pdf(form: ObservationForm) -> bytes:
    """Leichtgewichtiger PDF-Export ohne System-Abhängigkeiten."""
    pdf = FPDF(orientation="P", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "Hospitationsbogen – BLI 3.0", ln=True, align="C")

    pdf.set_font("Helvetica", size=11)

    def line(txt=""):
        pdf.multi_cell(0, 6, txt)

    # Meta
    line(f"Datum: {form.date}")
    line(f"Kolleg*in: {form.colleague}")
    line(f"Beobachter*in: {form.observer}")
    line(f"Fach/Klasse/Thema: {form.subject} / {form.grade} / {form.topic}")
    if form.school:
        line(f"Schule: {form.school}")
    if form.profile_focus:
        line("Profil-Fokus: " + ", ".join(form.profile_focus))
    pdf.ln(2)

    # Module + Kriterien
    for mk, mod in form.modules.items():
        pdf.set_font("Helvetica", "B", 13)
        line(f"{mk} – {BLI_DATA[mk]['title']}")
        pdf.set_font("Helvetica", size=11)
        for ck, cres in mod.criteria.items():
            line(f"{ck} {BLI_DATA[mk]['criteria'][ck]}")
            line(f"  Bewertung: {cres.rating}/4")
            if cres.comment:
                line(f"  Kommentar: {cres.comment}")
            pdf.ln(1)
        pdf.ln(2)

    # Stärken / nächste Schritte
    pdf.set_font("Helvetica", "B", 13)
    line("Stärken")
    pdf.set_font("Helvetica", size=11)
    line(form.strengths or "-")
    pdf.ln(2)

    pdf.set_font("Helvetica", "B", 13)
    line("Nächste Schritte (konkret, terminiert)")
    pdf.set_font("Helvetica", size=11)
    line(form.next_steps or "-")
    pdf.ln(2)

    # Scores
    per_module, overall = compute_scores(form)
    pdf.set_font("Helvetica", "B", 13)
    line("Zusammenfassung (Scores)")
    pdf.set_font("Helvetica", size=11)
    for mk, sc in per_module.items():
        line(f"{mk}: {sc:.2f} / 4")
    line(f"Gesamt (gewichtet): {overall:.2f} / 4")

    # Rückgabe als Bytes
    return pdf.output(dest="S").encode("latin-1", errors="ignore")

# ----------------------------- UI -----------------------------
st.title("Hospitationsbogen (BLI 3.0) – Generator")

# Profile Auswahl
with st.sidebar:
    st.header("Kolleg*in & Profil")
    profiles = st.session_state["profiles"]
    selected_profile = st.selectbox("Kolleg*in auswählen", list(profiles.keys()))
    new_name = ""
    if selected_profile == "— Neu —":
        new_name = st.text_input("Name für neues Profil", value="")
    focus_default = profiles[selected_profile]["focus"]
    weights_default = profiles[selected_profile]["weights"]

    st.subheader("Fokus-Merkmale")
    focus: List[str] = []
    cols = st.columns(4)
    all_modules = list(BLI_DATA.keys())
    for i, mk in enumerate(all_modules):
        with cols[i]:
            checked = st.checkbox(f"{mk}", value=(mk in focus_default))
            if checked:
                focus.append(mk)

    st.subheader("Gewichtungen (optional)")
    weights: Dict[str, float] = {}
    for mk in all_modules:
        weights[mk] = st.number_input(
            f"Gewicht {mk}",
            min_value=0.0, max_value=3.0, step=0.1,
            value=float(weights_default.get(mk, 1.0))
        )

    st.divider()
    st.subheader("Profil speichern/aktualisieren")
    if st.button("Profil speichern"):
        key = new_name.strip() if (selected_profile == "— Neu —" and new_name.strip()) else selected_profile
        st.session_state["profiles"][key] = {"focus": focus or focus_default, "weights": weights}
        st.success(f"Profil gespeichert: {key}")

# Meta
col1, col2, col3 = st.columns(3)
with col1:
    date = st.date_input("Datum", value=datetime.today())
    observer = st.text_input("Beobachter*in", value="")
with col2:
    colleague = st.text_input("Kolleg*in (Name)", value=(new_name if (selected_profile == "— Neu —" and new_name) else selected_profile))
    school = st.text_input("Schule (optional)", value="")
with col3:
    subject = st.text_input("Fach", value="")
    grade = st.text_input("Klasse/Jahrgang", value="")
topic = st.text_input("Thema/Sequenz", value="")

selected_modules = focus or list(BLI_DATA.keys())  # wenn kein Fokus gesetzt, dann alle

# Formularzustand initialisieren (bei Fokus-Wechsel neu erstellen)
if "form" not in st.session_state or st.session_state.get("form_modules") != selected_modules:
    st.session_state["form"] = init_form(selected_modules)
    st.session_state["form_modules"] = selected_modules

form: ObservationForm = st.session_state["form"]
form.date = date.strftime("%Y-%m-%d")
form.colleague = colleague
form.subject = subject
form.grade = grade
form.topic = topic
form.observer = observer
form.school = school
form.profile_focus = selected_modules
form.weights = weights

# Eingabe der Bewertungen
st.subheader("Bewertung je Kriterium (0–4)")
for mk in selected_modules:
    box = st.expander(f"{mk} – {BLI_DATA[mk]['title']}", expanded=False)
    with box:
        st.markdown("**Skala:** 0= nicht beobachtbar, 1= Ansatzweise, 2= Grundlegend, 3= Gut, 4= Sehr stark")
        for ck, ctext in BLI_DATA[mk]["criteria"].items():
            ccol = st.container()
            with ccol:
                c1, c2 = st.columns([1, 3])
                with c1:
                    rating_key = f"rating_{mk}_{ck}"
                    rating = st.slider(f"{mk}.{ck}:", 0, 4, value=form.modules[mk].criteria[ck].rating, key=rating_key)
                with c2:
                    # Auto-Kommentar vorschlagen
                    default_comment = AUTO_COMMENTS.get(rating, "")
                    comment_key = f"comment_{mk}_{ck}"
                    comment = st.text_area(f"Kommentar – {ctext}", value=form.modules[mk].criteria[ck].comment or default_comment, key=comment_key)
                # Speichern
                form.modules[mk].criteria[ck].rating = rating
                form.modules[mk].criteria[ck].comment = comment

# Stärken/Nächste Schritte
st.subheader("Zusammenfassung")
strengths = st.text_area("Stärken (konkret, evidenzbasiert)", value=form.strengths or "")
next_steps = st.text_area("Nächste Schritte (SMART formulieren)", value=form.next_steps or "")
form.strengths = strengths
form.next_steps = next_steps

# Live-Scores
if form.modules:
    per_module, overall = compute_scores(form)
    cols = st.columns(len(selected_modules) + 1)
    for i, mk in enumerate(selected_modules):
        with cols[i]:
            st.metric(f"{mk} Score", f"{per_module[mk]:.2f} / 4")
    with cols[-1]:
        st.metric("Gesamt (gewichtet)", f"{overall:.2f} / 4")

# Export
st.subheader("Exportieren")
c1, c2, c3 = st.columns(3)
with c1:
    docx_bytes = export_docx(form)
    fname = f"Hospitationsbogen_{form.colleague.replace(' ', '_')}_{form.date}.docx"
    st.download_button("DOCX herunterladen", data=docx_bytes, file_name=fname, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
with c2:
    pdf_bytes = export_pdf(form)
    pname = f"Hospitationsbogen_{form.colleague.replace(' ', '_')}_{form.date}.pdf"
    st.download_button("PDF herunterladen", data=pdf_bytes, file_name=pname, mime="application/pdf")
with c3:
    json_bytes = export_json(form)
    jname = f"Hospitationsbogen_{form.colleague.replace(' ', '_')}_{form.date}.json"
    st.download_button("JSON herunterladen", data=json_bytes, file_name=jname, mime="application/json")

st.info("Hinweis: Inhalte, Skalen und Gewichtungen sind schul-intern anpassbar. Fügen Sie bei Bedarf Indikatoren/Belege hinzu (z. B. Checklisten, Beobachtungsnotizen).")
